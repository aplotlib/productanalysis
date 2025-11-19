import streamlit as st
import pandas as pd
import plotly.express as px
import graphviz
from PIL import Image
import io
import re
import os
import json
import numpy as np
from datetime import datetime, timedelta
from docx import Document 
from docx.shared import Inches

# --- 1. SYSTEM CONFIGURATION ---
st.set_page_config(
    page_title="ORION | VIVE Health",
    page_icon="🧬",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- 2. UI THEME ---
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    :root { --primary: #0F172A; --accent: #2563EB; --bg-light: #F8FAFC; --card-bg: #FFFFFF; --border: #E2E8F0; }
    .stApp { background-color: var(--bg-light); font-family: 'Inter', sans-serif; color: #1E293B; }
    .brand-header { padding-bottom: 1rem; border-bottom: 1px solid var(--border); margin-bottom: 2rem; }
    .brand-title { font-size: 2.2rem; font-weight: 800; color: var(--primary); letter-spacing: -0.5px; margin: 0; }
    .brand-subtitle { font-size: 0.9rem; color: #64748B; font-weight: 500; margin-top: 0.25rem; }
    .feature-card { background: var(--card-bg); border: 1px solid var(--border); border-radius: 10px; padding: 24px; box-shadow: 0 2px 4px rgba(0,0,0,0.02); }
    div[data-testid="metric-container"] { background-color: white; border: 1px solid #E2E8F0; padding: 15px; border-radius: 8px; }
    .stTabs [data-baseweb="tab-list"] { gap: 8px; }
    .stTabs [data-baseweb="tab"] { height: 45px; background-color: white; border-radius: 6px; border: 1px solid #E2E8F0; padding: 0 20px; }
    .stTabs [aria-selected="true"] { background-color: var(--accent); color: white; border-color: var(--accent); }
    </style>
""", unsafe_allow_html=True)

# --- 3. INTELLIGENCE ENGINE (Gemini Fix) ---
class IntelligenceEngine:
    def __init__(self):
        self.client = None
        self.provider = None
        self.model_name = None
        self.available = False
        self.connection_error = None

    def _get_key(self, names, manual_key=None):
        if manual_key: return manual_key
        for name in names:
            if hasattr(st, "secrets") and name in st.secrets: return st.secrets[name]
            if os.environ.get(name): return os.environ.get(name)
        return None

    def configure_client(self, provider_choice, manual_key_input=None):
        self.available = False
        self.connection_error = None
        
        try:
            if "Gemini" in provider_choice:
                try:
                    import google.generativeai as genai
                except ImportError:
                    self.connection_error = "Missing `google-generativeai` library."
                    return

                api_key = self._get_key(["GOOGLE_API_KEY", "GEMINI_API_KEY"], manual_key_input)
                if api_key:
                    genai.configure(api_key=api_key)
                    self.client = genai
                    self.provider = "Google Gemini"
                    self.model_name = "gemini-1.5-pro" # Force latest Pro model
                    self.available = True
                else:
                    self.connection_error = "Missing Google API Key (Starts with AIza...)"

            elif "GPT" in provider_choice:
                try:
                    import openai
                except ImportError:
                    self.connection_error = "Missing `openai` library."
                    return

                api_key = self._get_key(["OPENAI_API_KEY"], manual_key_input)
                if api_key:
                    self.client = openai.OpenAI(api_key=api_key)
                    self.provider = "OpenAI"
                    self.model_name = "gpt-4o" if "4o" in provider_choice else "gpt-4o-mini"
                    self.available = True
                else:
                    self.connection_error = "Missing OpenAI API Key"
                    
        except Exception as e:
            self.connection_error = f"Connection Failed: {str(e)}"
            self.available = False

    def generate(self, prompt, temperature=0.3, json_mode=False):
        if not self.available: return None if json_mode else f"⚠️ AI Offline: {self.connection_error}"
        
        try:
            if "Gemini" in self.provider:
                model = self.client.GenerativeModel(self.model_name)
                config = {'temperature': temperature}
                if json_mode: config['response_mime_type'] = 'application/json'
                response = model.generate_content(prompt, generation_config=config)
                return response.text
                
            elif "OpenAI" in self.provider:
                kwargs = {
                    "model": self.model_name,
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": temperature
                }
                if json_mode: kwargs["response_format"] = {"type": "json_object"}
                response = self.client.chat.completions.create(**kwargs)
                return response.choices[0].message.content
        except Exception as e:
            return f"AI Error: {str(e)}"

    def analyze_image(self, image, prompt):
        if not self.available: return f"AI Offline: {self.connection_error}"
        try:
            if "Gemini" in self.provider:
                model = self.client.GenerativeModel(self.model_name)
                response = model.generate_content([prompt, image])
                return response.text
            elif "OpenAI" in self.provider:
                return "Vision requires Gemini 1.5 Pro."
        except Exception as e: return f"Vision Error: {e}"

    def generate_capa_draft(self, context):
        prompt = f"""
        Act as a Quality Engineer (ISO 13485). Create a CAPA investigation for:
        Product: {context.get('product')}
        Issue: {context.get('issue')}
        
        Return a JSON object with these keys: 
        issue_description, root_cause_analysis, corrective_action, preventive_action, effectiveness_plan.
        """
        res = self.generate(prompt, json_mode=True)
        try:
            if "```json" in res: res = res.split("```json")[1].split("```")[0]
            return json.loads(res)
        except: return None

if 'ai' not in st.session_state: st.session_state.ai = IntelligenceEngine()

# --- 4. DATA HANDLER (Safe Merge Logic) ---
class DataHandler:
    @staticmethod
    def clean_sku(sku):
        if pd.isna(sku): return "UNKNOWN"
        sku = str(sku).upper().strip()
        match = re.match(r"^([A-Z]+[0-9]+)", sku)
        return match.group(1) if match else sku

    @staticmethod
    def get_category(sku):
        if pd.isna(sku): return "Other"
        sku = str(sku).upper().strip()
        match = re.match(r"^([A-Z]+)", sku)
        return match.group(1) if match else "Other"

    @staticmethod
    def load_and_merge(files, period_days=90, manual_data=None):
        warnings = []
        cutoff_date = datetime.now() - timedelta(days=period_days)
        sales_records = []
        return_records = []

        # 1. Parse Files
        for file in files:
            try:
                fname = file.name.lower()
                raw = pd.read_excel(file, header=None)
                header_idx = 0
                for i in range(min(20, len(raw))):
                    row_str = raw.iloc[i].astype(str).str.lower().tolist()
                    if sum(1 for k in ['sku', 'product', 'asin', 'date', 'qty', 'sales'] if any(k in str(rs) for rs in row_str)) >= 2:
                        header_idx = i
                        break
                
                df = pd.read_excel(file, header=header_idx)
                df.columns = [str(c).strip() for c in df.columns]
                
                sku_col = next((c for c in df.columns if any(x in c.lower() for x in ['sku', 'product', 'default code'])), None)
                if not sku_col: continue

                if "odoo" in fname or "forecast" in fname:
                    sales_col = None
                    candidates = [c for c in df.columns if any(x in c.lower() for x in ['sales', 'qty', 'quantity', 'unnamed'])]
                    # Heuristic: Last numeric column in Odoo exports is usually the total forecast
                    for c in candidates[::-1]:
                         if pd.to_numeric(df[c], errors='coerce').sum() > 0:
                             sales_col = c
                             break
                    
                    if sales_col:
                        for _, row in df.iterrows():
                            sku = str(row[sku_col]).strip().upper()
                            qty = pd.to_numeric(row[sales_col], errors='coerce') or 0
                            sales_records.append({'SKU': sku, 'Sales': qty})

                elif "return" in fname or "pivot" in fname:
                    date_col = next((c for c in df.columns if 'date' in c.lower()), None)
                    reason_col = next((c for c in df.columns if any(x in c.lower() for x in ['reason', 'comment'])), None)
                    
                    if date_col:
                        df[date_col] = pd.to_datetime(df[date_col], errors='coerce')
                        df = df[df[date_col] >= cutoff_date]
                    
                    for _, row in df.iterrows():
                        sku = str(row[sku_col]).strip().upper()
                        issue = str(row[reason_col]) if reason_col and pd.notna(row[reason_col]) else None
                        return_records.append({'SKU': sku, 'Returns': 1, 'Issue': issue})

            except Exception as e: warnings.append(f"Error {file.name}: {str(e)}")

        # 2. Manual Data
        if manual_data is not None and not manual_data.empty:
            for _, row in manual_data.iterrows():
                sku = str(row.get('SKU', '')).strip().upper()
                if sku:
                    if 'Sales' in row and row['Sales'] > 0: sales_records.append({'SKU': sku, 'Sales': row['Sales']})
                    if 'Returns' in row and row['Returns'] > 0:
                        for _ in range(int(row['Returns'])): return_records.append({'SKU': sku, 'Returns': 1, 'Issue': 'Manual Entry'})

        # 3. Aggregate
        sales_df = pd.DataFrame(sales_records).groupby('SKU')['Sales'].sum().reset_index() if sales_records else pd.DataFrame(columns=['SKU', 'Sales'])
        returns_df = pd.DataFrame(return_records)
        
        if not returns_df.empty:
            returns_agg = returns_df.groupby('SKU')['Returns'].sum().reset_index()
            # Aggregate issues (limit to top 5 for sanity)
            issues_agg = returns_df.groupby('SKU')['Issue'].apply(lambda x: list(filter(None, x))[:5]).reset_index()
            returns_final = pd.merge(returns_agg, issues_agg, on='SKU')
        else:
            returns_final = pd.DataFrame(columns=['SKU', 'Returns', 'Issue'])

        master = pd.merge(sales_df, returns_final, on='SKU', how='outer').fillna(0)
        
        # 4. Final Polish (Prevent KeyErrors)
        if not master.empty:
            master['Returns'] = master['Returns'].astype(int)
            master['Sales'] = master['Sales'].astype(int)
            master['Parent_SKU'] = master['SKU'].apply(DataHandler.clean_sku)
            master['Category'] = master['SKU'].apply(DataHandler.get_category)
            master['Return_Rate'] = (master['Returns'] / master['Sales'] * 100).fillna(0)
            master['Return_Rate'] = master.apply(lambda x: 100.0 if x['Sales'] == 0 and x['Returns'] > 0 else x['Return_Rate'], axis=1)
            if 'Issue' in master.columns:
                master['Issues'] = master['Issue'].apply(lambda x: x if isinstance(x, list) else [])
                master = master.drop(columns=['Issue'])
            else:
                master['Issues'] = [[] for _ in range(len(master))]
        else:
            # Return empty schema to prevent UI crashes
            master = pd.DataFrame(columns=['SKU', 'Sales', 'Returns', 'Parent_SKU', 'Category', 'Return_Rate', 'Issues'])

        return master, warnings

# --- 5. EXPORT UTILS (Word Doc) ---
def create_word_doc(capa_data):
    doc = Document()
    doc.add_heading(f"CAPA Report: {capa_data.get('id', 'New')}", 0)
    
    doc.add_heading('1. Initiation', level=1)
    doc.add_paragraph(f"Product SKU: {capa_data.get('sku', 'N/A')}")
    doc.add_paragraph(f"Issue Description: {capa_data.get('desc', 'N/A')}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")

    doc.add_heading('2. Investigation (Root Cause)', level=1)
    doc.add_paragraph(capa_data.get('root_cause_analysis', 'Pending investigation.'))
    
    if 'fishbone' in capa_data and capa_data['fishbone']:
        doc.add_heading('Fishbone Analysis:', level=2)
        for cat, cause in capa_data['fishbone']:
            doc.add_paragraph(f"{cat}: {cause}", style='List Bullet')

    doc.add_heading('3. Risk Assessment (FMEA)', level=1)
    if 'risks' in capa_data and not capa_data['risks'].empty:
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ['Failure Mode', 'Effect', 'Sev', 'Occ', 'Det', 'RPN']
        for i, h in enumerate(headers): hdr_cells[i].text = h
        
        for _, row in capa_data['risks'].iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row.get('Failure Mode', ''))
            row_cells[1].text = str(row.get('Effect', ''))
            row_cells[2].text = str(row.get('Sev', ''))
            row_cells[3].text = str(row.get('Occ', ''))
            row_cells[4].text = str(row.get('Det', ''))
            row_cells[5].text = str(row.get('RPN', ''))

    doc.add_heading('4. Action Plan', level=1)
    doc.add_paragraph(f"Corrective Action: {capa_data.get('corrective_action', 'Pending')}")
    doc.add_paragraph(f"Preventive Action: {capa_data.get('preventive_action', 'Pending')}")
    doc.add_paragraph(f"Verification Plan: {capa_data.get('effectiveness_plan', 'Pending')}")

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- 6. UI MODULES ---

def render_sidebar():
    with st.sidebar:
        st.markdown("### System Controls")
        
        model_choice = st.selectbox("Provider", 
            ["Google Gemini 1.5 Pro", "Google Gemini 1.5 Flash", "OpenAI GPT-4o"], index=0)
        
        manual_key = None
        needs_key = ("Gemini" in model_choice and "GOOGLE_API_KEY" not in st.secrets) or \
                    ("GPT" in model_choice and "OPENAI_API_KEY" not in st.secrets)
        
        if needs_key:
            st.warning(f"No secret found for {model_choice[:10]}...")
            manual_key = st.text_input("Enter API Key", type="password")
        
        st.session_state.ai.configure_client(model_choice, manual_key)
        
        if st.session_state.ai.available:
            st.markdown(f"<span style='color:green'>● {model_choice} Online</span>", unsafe_allow_html=True)
        else:
            st.markdown(f"<span style='color:red'>● Offline</span>", unsafe_allow_html=True)
            if st.session_state.ai.connection_error:
                st.caption(f"Error: {st.session_state.ai.connection_error}")

        st.markdown("---")
        st.session_state.period_days = st.selectbox("Analysis Period", [30, 60, 90, 180, 365], index=2)
        return st.radio("Navigation", ["Dashboard", "Data Ingestion", "Vision", "CAPA Manager", "Strategy"], label_visibility="collapsed")

def render_ingestion():
    st.markdown("### Data Ingestion")
    t1, t2 = st.tabs(["File Upload", "Manual Entry"])
    
    with t1:
        st.markdown("""<div class='feature-card'>Upload <b>Odoo Sales</b> and <b>Return Reports</b>.</div>""", unsafe_allow_html=True)
        files = st.file_uploader("Upload Files", accept_multiple_files=True)
        
    with t2:
        if 'manual_entry_df' not in st.session_state:
            st.session_state.manual_entry_df = pd.DataFrame([{"SKU": "", "Sales": 0, "Returns": 0}])
        manual_data = st.data_editor(st.session_state.manual_entry_df, num_rows="dynamic", use_container_width=True)
        st.session_state.manual_entry_df = manual_data

    if st.button("Process All Data", type="primary"):
        with st.spinner("Merging..."):
            df, warns = DataHandler.load_and_merge(files if files else [], st.session_state.period_days, manual_data)
            st.session_state.master_data = df
            st.session_state.filtered_data = df
            for w in warns: st.warning(w)
            st.success(f"Processed {len(df)} SKUs.")

def render_dashboard():
    st.markdown("### Executive Dashboard")
    if 'filtered_data' not in st.session_state or st.session_state.filtered_data.empty:
        st.info("No data loaded.")
        return
        
    df = st.session_state.filtered_data
    
    c1, c2, c3 = st.columns(3)
    cats = ["All"] + sorted(df['Category'].unique().tolist())
    sel_cat = c1.selectbox("Category", cats)
    if sel_cat != "All": df = df[df['Category'] == sel_cat]
    
    parents = ["All"] + sorted(df['Parent_SKU'].unique().tolist())
    sel_parent = c2.selectbox("Parent SKU", parents)
    if sel_parent != "All": df = df[df['Parent_SKU'] == sel_parent]
    
    total_sales = int(df['Sales'].sum())
    total_returns = int(df['Returns'].sum())
    avg_rate = (total_returns / total_sales * 100) if total_sales > 0 else 0
    
    m1, m2, m3 = st.columns(3)
    m1.metric("Total Sales", f"{total_sales:,}")
    m2.metric("Total Returns", f"{total_returns:,}")
    m3.metric("Avg Return Rate", f"{avg_rate:.2f}%")
    
    st.markdown("---")
    if not df.empty:
        chart_df = df.groupby('Parent_SKU')[['Sales', 'Returns']].sum().reset_index()
        chart_df['Rate'] = (chart_df['Returns'] / chart_df['Sales'] * 100).fillna(0)
        fig = px.bar(chart_df.sort_values('Rate', ascending=False).head(15), x='Parent_SKU', y=['Sales', 'Returns'], title="Top Return Drivers")
        st.plotly_chart(fig, use_container_width=True)

def render_capa():
    st.markdown("### CAPA Manager")
    if 'capa_data' not in st.session_state:
        st.session_state.capa_data = {
            'id': f"CAPA-{datetime.now().strftime('%y%m%d')}",
            'risks': pd.DataFrame([{"Failure Mode": "Example", "Effect": "Return", "Sev": 1, "Occ": 1, "Det": 1, "RPN": 1}])
        }
    data = st.session_state.capa_data
    
    if 'capa_prefill' in st.session_state:
        data['desc'] = st.session_state.capa_prefill.get('desc')
        del st.session_state.capa_prefill

    t1, t2, t3, t4 = st.tabs(["Intake", "Investigation", "FMEA Risk", "Action & Export"])

    with t1:
        c1, c2 = st.columns(2)
        data['id'] = c1.text_input("CAPA ID", data['id'])
        data['sku'] = c2.text_input("SKU", data.get('sku',''))
        data['desc'] = st.text_area("Issue Description", data.get('desc',''))
        if st.button("✨ AI Auto-Draft"):
            with st.spinner("Drafting..."):
                draft = st.session_state.ai.generate_capa_draft({'product':data['sku'], 'issue':data['desc']})
                if draft: 
                    data.update(draft)
                    st.success("Draft populated!")
                    st.rerun()

    with t2:
        st.markdown("#### Root Cause Analysis")
        data['root_cause_analysis'] = st.text_area("5 Whys / Fishbone Analysis", value=data.get('root_cause_analysis', ''), height=150)
        
        c_fish1, c_fish2 = st.columns([1, 3])
        with c_fish1:
            cat = st.selectbox("Category", ["Man", "Machine", "Material", "Method", "Env"])
            cause = st.text_input("Cause")
            if st.button("Add Cause"):
                if 'fishbone' not in data: data['fishbone'] = []
                data['fishbone'].append((cat, cause))
        with c_fish2:
            if data.get('fishbone'):
                g = graphviz.Digraph()
                g.attr(rankdir='LR')
                g.node('Problem', 'Issue')
                for c, r in data['fishbone']:
                    g.edge(c, 'Problem')
                    g.edge(r, c)
                st.graphviz_chart(g)

    with t3:
        st.markdown("#### Risk Analysis (FMEA)")
        st.caption("Edit table directly. RPN auto-calculates on save.")
        
        edited_df = st.data_editor(
            data['risks'],
            num_rows="dynamic",
            column_config={
                "Sev": st.column_config.NumberColumn("Severity", min_value=1, max_value=10, step=1),
                "Occ": st.column_config.NumberColumn("Occurrence", min_value=1, max_value=10, step=1),
                "Det": st.column_config.NumberColumn("Detection", min_value=1, max_value=10, step=1),
                "RPN": st.column_config.NumberColumn("RPN", disabled=True),
            },
            use_container_width=True,
            key="fmea_editor"
        )
        # Auto-calc RPN
        edited_df['RPN'] = edited_df['Sev'] * edited_df['Occ'] * edited_df['Det']
        data['risks'] = edited_df # Persist state

    with t4:
        st.markdown("#### Action Plan")
        data['corrective_action'] = st.text_area("Corrective Action", value=data.get('corrective_action', ''))
        data['preventive_action'] = st.text_area("Preventive Action", value=data.get('preventive_action', ''))
        data['effectiveness_plan'] = st.text_area("Effectiveness Verification", value=data.get('effectiveness_plan', ''))
        
        st.markdown("---")
        doc_bytes = create_word_doc(data)
        st.download_button(
            label="📄 Download Report (.docx)",
            data=doc_bytes,
            file_name=f"{data['id']}_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

def render_vision():
    st.markdown("### Vision Diagnostics")
    img = st.file_uploader("Upload Image", type=['png','jpg'])
    if img:
        st.image(img, width=300)
        if st.button("Analyze"):
            res = st.session_state.ai.analyze_image(Image.open(img), "Analyze defect severity and potential root cause.")
            st.session_state.vision_res = res
    if 'vision_res' in st.session_state:
        st.markdown(f"<div class='feature-card'>{st.session_state.vision_res}</div>", unsafe_allow_html=True)
        if st.button("Escalate to CAPA"):
            st.session_state.capa_prefill = {'desc': st.session_state.vision_res}
            st.success("Sent to CAPA")

def render_strategy():
    st.markdown("### Strategy Architect")
    st.info("Generate SOPs and Work Instructions specific to your problem context.")
    
    doc_type = st.selectbox("Document Type", ["SOP (Standard Operating Procedure)", "Work Instruction", "Quality Policy"])
    
    with st.form("strat_form"):
        c1, c2 = st.columns(2)
        problem = c1.text_input("What problem are we solving?")
        solution = c2.text_input("What is the best known way to solve it?")
        
        c3, c4 = st.columns(2)
        stakeholders = c3.text_input("Who are the stakeholders?")
        limitations = c4.text_input("Are there limitations or constraints?")
        
        detail = st.select_slider("Thoroughness Level", options=["Brief", "Standard", "Detailed/Audit-Ready"])
        
        if st.form_submit_button("Generate Document"):
            with st.spinner("Architecting..."):
                prompt = f"""
                Write a {doc_type} for a Medical Device QMS.
                Context:
                - Problem: {problem}
                - Solution Method: {solution}
                - Stakeholders: {stakeholders}
                - Limitations: {limitations}
                - Detail Level: {detail}
                
                Format nicely with Markdown headers, Purpose, Scope, and Procedure steps.
                """
                res = st.session_state.ai.generate(prompt)
                st.session_state.strat_res = res

    if 'strat_res' in st.session_state:
        st.markdown("---")
        st.markdown(st.session_state.strat_res)
        st.download_button("Download .MD", st.session_state.strat_res, "document.md")

# --- MAIN ---
def main():
    nav = render_sidebar()
    if nav == "Dashboard": render_dashboard()
    elif nav == "Data Ingestion": render_ingestion()
    elif nav == "CAPA Manager": render_capa()
    elif nav == "Vision": render_vision()
    elif nav == "Strategy": render_strategy()

if __name__ == "__main__":
    main()
